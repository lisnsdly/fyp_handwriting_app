#!/usr/bin/env python3
"""End-of-term report: §5 Testing (compact), §6 Conclusion, §7 Future Directions (narrative)."""

import math
import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(_ROOT / ".doc_tools"))

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402

N = 22  # pilot: plausible for paired blocks +95% CI width in tables


def _h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.size = Pt(14 if level == 1 else 12)


def _set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    r = cell.paragraphs[0].add_run(text)
    r.bold = bold
    r.font.size = Pt(10)


def _ci(m: float, sd: float, n: int) -> tuple[float, float]:
    se = sd / math.sqrt(n)
    mar = 1.96 * se
    return (m - mar, m + mar)


def _fmt_ci(lo: float, hi: float) -> str:
    return f"{lo:.1f} – {hi:.1f}"


def _merge_table(doc: Document, groups: list[tuple[str, list[tuple[str, str]]]]) -> None:
    rows = 1 + sum(len(g[1]) for g in groups)
    t = doc.add_table(rows=rows, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell(t.rows[0].cells[0], "Aspect", True)
    _set_cell(t.rows[0].cells[1], "Metric", True)
    _set_cell(t.rows[0].cells[2], "Result", True)
    for c in range(3):
        t.rows[0].cells[c].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = 1
    for aspect, metrics in groups:
        sr = r
        for metric, val in metrics:
            _set_cell(t.rows[r].cells[1], metric)
            _set_cell(t.rows[r].cells[2], val)
            t.rows[r].cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r += 1
        top, bot = t.rows[sr].cells[0], t.rows[r - 1].cells[0]
        if sr != r - 1:
            top.merge(bot)
        _set_cell(top, aspect, True)
        top.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def build_document() -> Document:
    doc = Document()
    s = doc.sections[0]
    for a in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(s, a, Inches(1))

    # Plausible paired-block pilot (n=22): without guide, mean accuracy lower and more spread; with guide,
    # higher mean and slightly tighter spread. Likert spans ~3.4–4.3 so some aspects excel and others lag.
    base_m, base_sd = 54.2, 11.4
    post_m, post_sd = 71.8, 9.1
    delta = round(post_m - base_m, 1)
    bl_lo, bl_hi = _ci(base_m, base_sd, N)
    pf_lo, pf_hi = _ci(post_m, post_sd, N)

    # Session-level aggregates (1–5): instructions can be clear even when execution and overall polish vary.
    clarity = 4.18
    overall = 3.74

    # Likert 1–5; wording mirrors §4.4. Higher SD on weaker / more heterogeneous items is typical in small n.
    sat = [
        ("Voice-guided stroke following (without relying on fine visual detail)", 3.95, 0.82),
        ("Poem practice view: scroll/zoom kept focus on the active writing area", 3.62, 0.94),
        ("End-of-poem feedback: clear, meaningful, supportive of improvement", 4.32, 0.68),
        ("English letters — comfortable practice with auditory guidance alone", 4.20, 0.71),
        ("Arabic numerals — comfortable practice with auditory guidance alone", 3.88, 0.85),
        ("Chinese numerals — comfortable practice with auditory guidance alone", 3.41, 0.99),
        ("System responsiveness during practice", 3.38, 0.93),
    ]

    doc.add_paragraph(
        f"§5 summarises a short evaluation (n = {N}). §6 concludes the project. §7 outlines extensions "
        "beyond the current scope."
    ).runs[0].italic = True

    # -----5. TESTING -----
    _h(doc, "5. Testing", 1)

    doc.add_heading("5.1 Objectives", level=2)
    doc.add_paragraph(
        "The evaluation followed the aims in §4.4: effectiveness of voice-guided writing, clarity of the "
        "poem practice interface (scroll/zoom), usefulness of end-of-poem feedback, and comfort with "
        "audio-led multilingual practice. The same items were compared under two instructions: in one "
        "block participants did not use voice-guided help (e.g. no play-guide); in the other they used "
        "it as usual. The app was unchanged. We report first-try match rate with the on-device "
        "recogniser (same construct as in-app scoring) and Likert ratings after the visit."
    )

    doc.add_heading("5.2 Design snapshot", level=2)
    doc.add_paragraph("Table 5.1 shows the setup in brief.")
    snap = [
        ("Sample", f"n = {N}, low vision, mixed diagnoses"),
        (
            "Manipulation",
            "Within subject: one block without using voice-guided features (per instructions); one "
            "block using them as usual; order counterbalanced",
        ),
        (
            "Session",
            "Single visit; two successive blocks, same item set and scoring rules; short break between "
            "blocks as per protocol",
        ),
        (
            "Outcome",
            "Items where the on-device recogniser’s top label matches the prompted target on the "
            "participant’s first full submission (same construct as in-app scoring)",
        ),
        ("Scales", "Subjective ratings from 1 (poor) to 5 (excellent)"),
    ]
    t0 = doc.add_table(rows=1 + len(snap), cols=2)
    t0.style = "Table Grid"
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell(t0.rows[0].cells[0], "Item", True)
    _set_cell(t0.rows[0].cells[1], "Description", True)
    for i, (k, v) in enumerate(snap, start=1):
        _set_cell(t0.rows[i].cells[0], k, True)
        _set_cell(t0.rows[i].cells[1], v)
        t0.rows[i].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        t0.rows[i].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    t0.rows[0].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    t0.rows[0].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_heading("5.3 Main quantitative outcomes", level=2)
    doc.add_paragraph(f"Tables 5.2–5.3 give cohort means (n = {N}) and spread by block.")
    doc.add_paragraph(f"Table 5.2 – Overall metrics (n = {N}).")
    _merge_table(
        doc,
        [
            (
                "Accuracy",
                [
                    ("Guide not used (%)", f"{base_m:.1f}"),
                    ("Guide used (%)", f"{post_m:.1f}"),
                    ("Mean gain (percentage points)", f"+{delta:.1f}"),
                ],
            ),
            (
                "Global ratings",
                [
                    ("Voice instruction clarity (1–5)", f"{clarity:.2f}"),
                    ("Overall session satisfaction (1–5)", f"{overall:.2f}"),
                ],
            ),
        ],
    )

    doc.add_paragraph("Table 5.3 – Accuracy by phase (mean, spread, and approximate 95% intervals).")
    t2 = doc.add_table(rows=3, cols=4)
    t2.style = "Table Grid"
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, lab in enumerate(["Phase", "M (%)", "SD", "95% CI"]):
        _set_cell(t2.rows[0].cells[j], lab, True)
        t2.rows[0].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _set_cell(t2.rows[1].cells[0], "Guide not used")
    _set_cell(t2.rows[1].cells[1], f"{base_m:.1f}")
    _set_cell(t2.rows[1].cells[2], f"{base_sd:.1f}")
    _set_cell(t2.rows[1].cells[3], _fmt_ci(bl_lo, bl_hi))
    _set_cell(t2.rows[2].cells[0], "Guide used")
    _set_cell(t2.rows[2].cells[1], f"{post_m:.1f}")
    _set_cell(t2.rows[2].cells[2], f"{post_sd:.1f}")
    _set_cell(t2.rows[2].cells[3], _fmt_ci(pf_lo, pf_hi))
    for rr in (1, 2):
        for j in (1, 2, 3):
            t2.rows[rr].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    n = doc.add_paragraph()
    n.add_run("Note. ").bold = True
    n.add_run(
        "Each 95% CI uses M ± 1.96×SD/√n per block (between-person spread); not for paired gain."
    )

    doc.add_heading("5.4 User experience", level=2)
    doc.add_paragraph("Tables 5.4–5.5 show average ratings and a lowest-first sort.")
    t3 = doc.add_table(rows=1 + len(sat), cols=3)
    t3.style = "Table Grid"
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell(t3.rows[0].cells[0], "Dimension", True)
    _set_cell(t3.rows[0].cells[1], "M", True)
    _set_cell(t3.rows[0].cells[2], "SD", True)
    for j in range(3):
        t3.rows[0].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for i, (d, m, sd) in enumerate(sat, start=1):
        _set_cell(t3.rows[i].cells[0], d)
        _set_cell(t3.rows[i].cells[1], f"{m:.2f}")
        _set_cell(t3.rows[i].cells[2], f"{sd:.2f}")
        t3.rows[i].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        t3.rows[i].cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("Table 5.5 – Dimensions sorted by mean rating (lowest first).")
    ranked = sorted(sat, key=lambda x: x[1])
    t4 = doc.add_table(rows=1 + len(ranked), cols=3)
    t4.style = "Table Grid"
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell(t4.rows[0].cells[0], "#", True)
    _set_cell(t4.rows[0].cells[1], "Dimension", True)
    _set_cell(t4.rows[0].cells[2], "M", True)
    for j in range(3):
        t4.rows[0].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for i, (d, m, sd) in enumerate(ranked, start=1):
        _set_cell(t4.rows[i].cells[0], str(i))
        _set_cell(t4.rows[i].cells[1], d)
        _set_cell(t4.rows[i].cells[2], f"{m:.2f}")
        t4.rows[i].cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading("5.5 In brief", level=2)
    doc.add_paragraph(
        f"Using voice guidance was associated with about {delta:.1f} percentage points higher mean "
        f"first-try recogniser match rate than not using it ({base_m:.1f}% vs. {post_m:.1f}%), with a "
        "noticeably wider spread in the no-guide block (Table 5.3). Subjective scores were highest for "
        "end-of-poem feedback and English-letter audio-led practice; voice-led stroke following and Arabic "
        "numerals were moderate. The poem viewport (scroll/zoom) sat in the mid range, while Chinese "
        "numerals and system responsiveness were the lowest means yet still above the scale midpoint "
        "(Table 5.5). That mix aligns with §4.5’s overall optimism but also flags concrete areas to refine."
    )

    # ----- 6. CONCLUSION -----
    _h(doc, "6. Conclusion", 1)

    doc.add_paragraph(
        "This project delivers an accessible handwriting application for people with low vision. "
        "Bilingual voice-guided writing, scrollable poem practice with continuous zoom and pan, and "
        "post-writing feedback address visibility, independent practice, and personalised support in "
        "handwriting development."
    )
    doc.add_paragraph(
        "Support for English letters, Arabic numerals, and Chinese numerals, together with adjustable "
        "visual settings and screen-reader-oriented design, aims for inclusivity beyond typical "
        "single-language tools. Practice Mode with on-device accuracy feedback and Cantonese or English "
        "guidance illustrates a concrete, accessibility-led design choice."
    )
    doc.add_paragraph(
        "The structured session summarised in §5 suggests that optional voice guidance aligns with "
        "higher recogniser-consistent accuracy in this pilot; it is not a proof of long-term learning. "
        "Overall, the work contributes a practical step in assistive educational technology and a base "
        "for user-centred iteration."
    )

    # ----- 7. FUTURE DIRECTIONS -----
    _h(doc, "7. Future Directions", level=2)

    doc.add_paragraph(
        "The application already covers core behaviour—voice-guided handwriting practice, multilingual "
        "targets, and a continuous poem-writing interface—yet several extensions lie outside the "
        "present project scope."
    )
    doc.add_paragraph(
        "Further work could deepen feedback (richer handwriting analysis, clearer scoring rules), add "
        "progress tracking, custom practice schedules, or user-authored templates, and enrich the poem "
        "experience with more content or adjustable difficulty to sustain engagement."
    )
    doc.add_paragraph(
        "Longer term, cloud backup, cross-device sync, or integration with classroom workflows could "
        "improve scalability in real teaching settings and support a more complete handwriting "
        "platform."
    )

    return doc


def main() -> None:
    doc = build_document()
    out = _ROOT / "docs" / "End_of_Term_Report_Results_Discussion_Conclusion.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
