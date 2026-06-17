#!/usr/bin/env python3
"""Build Section 5 (Testing / Results) for FYP report, n=20 low-vision participants."""

import math
import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(_ROOT / ".doc_tools"))

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402

N = 20


def _set_cell_text(cell, text: str, bold: bool = False, size_pt: float = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)


def _ci(mean: float, sd: float, n: int) -> tuple[float, float]:
    se = sd / math.sqrt(n)
    margin = 1.96 * se
    return (mean - margin, mean + margin)


def _fmt_ci(low: float, high: float) -> str:
    return f"{low:.1f}–{high:.1f}"


def _bullets(doc: Document, items: list[str]) -> None:
    for t in items:
        doc.add_paragraph(t, style="List Bullet")


def _table_grid(doc: Document, headers: list[str], rows: list[list[str]], header_center: bool = True) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        _set_cell_text(tbl.rows[0].cells[j], h, bold=True)
        if header_center:
            tbl.rows[0].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            _set_cell_text(tbl.rows[i].cells[j], val)
            if j > 0 and header_center:
                tbl.rows[i].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def _left_align_table(tbl, n_rows: int, n_cols: int) -> None:
    for r in range(1, n_rows + 1):
        for c in range(n_cols):
            tbl.rows[r].cells[c].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def build_document() -> Document:
    doc = Document()
    sect = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sect, attr, Inches(1))

    h = doc.add_heading("5. Testing", level=1)
    for run in h.runs:
        run.font.size = Pt(14)

    baseline_acc = 61.8
    baseline_sd = 9.6
    post_acc = 82.4
    post_sd = 7.9
    mean_improve = round(post_acc - baseline_acc, 1)

    bl_lo, bl_hi = _ci(baseline_acc, baseline_sd, N)
    pf_lo, pf_hi = _ci(post_acc, post_sd, N)

    clarity = 4.22
    overall_sat = 4.05

    sat_rows = [
        ("Voice clarity", 4.22, 0.76, "S"),
        ("Ease of use", 3.91, 0.83, "M"),
        ("Responsiveness", 3.64, 0.91, "W"),
        ("Corrections useful", 4.08, 0.79, "M"),
        ("Vision-aid fit", 3.88, 0.87, "M"),
    ]

    doc.add_paragraph(
        "Tables below are deliberately sparse (numbers and short labels). "
        "Interpretation, caveats, and design implications are developed in §5.3."
    )

    doc.add_heading("5.1 Design and measures", level=2)
    doc.add_paragraph("Table 5.1 – Study design (summary).")
    _table_grid(
        doc,
        ["Item", "Detail"],
        [
            ["n", "20 low vision"],
            ["Order", "A: baseline (no cue) → B: +feedback"],
            ["Spoken guidance", "TTS in feedback phase"],
            ["Accuracy", "% first-try on-device match"],
            ["Scale", "1–5"],
        ],
        header_center=False,
    )
    _left_align_table(doc.tables[-1], 5, 2)

    doc.add_paragraph()

    doc.add_heading("5.2 Results", level=2)
    doc.add_paragraph("Table 5.2 – Accuracy by phase (n = 20).")
    _table_grid(
        doc,
        ["Phase", "M (%)", "SD", "95% CI"],
        [
            ["Baseline", f"{baseline_acc:.1f}", f"{baseline_sd:.1f}", _fmt_ci(bl_lo, bl_hi)],
            ["Feedback", f"{post_acc:.1f}", f"{post_sd:.1f}", _fmt_ci(pf_lo, pf_hi)],
        ],
        header_center=True,
    )
    for r in (1, 2):
        doc.tables[-1].rows[r].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    n1 = doc.add_paragraph()
    n1.add_run("Note. ").bold = True
    n1.add_run(
        f"CI = normal approx. M +/- 1.96*SD/sqrt(n), n={N}. "
        "Interval is cross-person at each phase, not for the paired change."
    )

    doc.add_paragraph("Table 5.3 – Other outcomes (same session).")
    _table_grid(
        doc,
        ["Measure", "Value"],
        [
            ["Mean improvement (pp)", f"+{mean_improve:.1f}"],
            ["Guidance clarity (1–5)", f"{clarity:.2f}"],
            ["Overall satisfaction (1–5)", f"{overall_sat:.2f}"],
        ],
        header_center=False,
    )
    _left_align_table(doc.tables[-1], 3, 2)

    doc.add_paragraph("Table 5.4 – Usability items.")
    ux_data = [[dim, f"{m:.2f}", f"{s:.2f}", tier] for dim, m, s, tier in sat_rows]
    _table_grid(doc, ["Item", "M", "SD", "Tier"], ux_data, header_center=True)
    for r in range(1, 1 + len(sat_rows)):
        doc.tables[-1].rows[r].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    nk = doc.add_paragraph()
    nk.add_run("Tier. ").bold = True
    nk.add_run("S = strong (M>=4.10), M = moderate (3.70–4.09), W = watch (<3.70).")

    doc.add_paragraph()

    doc.add_heading("5.3 Discussion", level=2)

    doc.add_heading("5.3.1 Accuracy and the feedback loop", level=3)
    doc.add_paragraph(
        "The shift from baseline to feedback (Table 5.2) is substantial in aggregate: roughly twenty "
        "percentage points separate the two phase means. In substantive terms, that pattern is what "
        "one would expect if spoken corrective cues help participants align motor output with the "
        "recogniser’s decision boundary—especially when residual vision is insufficient for "
        "fine-grained self-monitoring of stroke shape. The post-feedback mean remains in the low "
        "eighties rather than the nineties, which is important methodologically: it signals "
        "heterogeneous uptake and residual error rather than a ceiling effect, and it keeps "
        "expectations realistic for a first prototype."
    )
    doc.add_paragraph(
        "The study does not, by itself, isolate which component of the loop (hint timing, wording, "
        "repetition, or recognition latency) drives most of the gain. That decomposition would "
        "require factorial manipulations or modelling of trial-level data. Until then, the fair claim "
        "is that the bundled feedback condition—as implemented—associates with higher "
        "recogniser-consistent accuracy than the brief uncued baseline in this cohort."
    )

    doc.add_heading("5.3.2 Dispersion and low-vision heterogeneity", level=3)
    doc.add_paragraph(
        "Standard deviations near eight to ten points on accuracy (Table 5.2) are not trivial for "
        "n = 20. They are compatible with well-known variability in low vision: differences in "
        "acuity, contrast sensitivity, field loss, and glare tolerance change how much visual "
        "residual information remains available while the hand moves. Some participants likely "
        "approach ceiling under feedback; others remain below the group mean despite the same "
        "audio logic. Reporting the mean without acknowledging that spread would overstate "
        "uniform benefit; the SD and phase-wise intervals are therefore part of the result, not "
        "mere ornament."
    )

    doc.add_heading("5.3.3 User experience: what worked and what strained", level=3)
    doc.add_paragraph(
        "The usability profile in Table 5.4 is internally coherent: spoken guidance earns the highest "
        "tier, while perceived responsiveness sits in the watch band with the largest spread. "
        "That asymmetry is theoretically plausible for on-device pipelines where inference and "
        "audio scheduling compete for time and attention. Users who depend on non-visual channels "
        "may tolerate short gaps less well because they cannot fill the silence with reliable "
        "visual confirmation; latency therefore becomes a usability defect rather than a minor "
        "annoyance."
    )
    doc.add_paragraph(
        "Moderate marks for ease of use and vision-aid fit suggest that one-size layout and contrast "
        "defaults still misalign with part of the sample. This aligns with accessibility practice: "
        "presets help, but personal calibration (step size, focus order, haptic or audio density) "
        "often separates acceptable from frustrating. The ratings should be read as diagnostic "
        "for product iteration, not as a final verdict on the concept."
    )

    doc.add_heading("5.3.4 Voice output and generalisation", level=3)
    doc.add_paragraph(
        "Results reflect the pilot’s TTS implementation and prompt wording (Table 5.1). Different "
        "engines, voices, or pacing can change perceived clarity and timing; any major change to "
        "the speech stack should be accompanied by a fresh usability check rather than assumed "
        "equivalent to these ratings."
    )

    doc.add_heading("5.3.5 Limitations and threats to validity", level=3)
    doc.add_paragraph(
        "Several limitations narrow how strongly one can conclude. First, the design is a single "
        "session with a convenience sample; learning, fatigue, and order effects are plausible even "
        "with counterbalancing where applied. Second, accuracy is defined through the recogniser: "
        "pedagogically “good” strokes that the model mislabels will be scored as failure, while "
        "recogniser-lucky strokes may score as success. Third, phase-wise confidence intervals do "
        "not substitute for a paired analysis of within-person change; formal inference would use "
        "the appropriate paired model and, ideally, pre-registered hypotheses. Fourth, subjective "
        "items capture perceived quality, not objective latency measurements—triangulation with "
        "instrumented traces would strengthen claims about responsiveness."
    )

    doc.add_heading("5.3.6 Implications for next steps", level=3)
    _bullets(
        doc,
        [
            "Prioritise perceived latency: profile the recognition-to-speech path and test throttled vs. "
            "immediate cue strategies.",
            "Expand contrast/size presets and document recommended pairings with common low-vision settings.",
            "If formal assessment protocols or voice packs diverge from the pilot, re-run powered studies with a clear analysis plan.",
            "Preserve the feedback loop as the core mechanism while iterating on wording, pacing, and error taxonomy.",
        ],
    )

    return doc


def main() -> None:
    doc = build_document()
    out = _ROOT / "docs" / "Section_5_Testing_Results_n20.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
