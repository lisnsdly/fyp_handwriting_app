#!/usr/bin/env python3
"""Build Section 4.2 as a Word document with a feature-vs-competitors table."""

import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(_ROOT / ".doc_tools"))

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402

# Table cells: strong / partial / weak — scannable at a glance.
_SYM = {"ok": "\u2713", "mid": "~", "no": "\u2014"}  # check, tilde, em dash
_LEGEND = (
    f"Legend: {_SYM['ok']} = strong support · {_SYM['mid']} = partial or varies · "
    f"{_SYM['no']} = weak, absent, or N/A for that approach."
)


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def build_document() -> Document:
    doc = Document()
    sect = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sect, attr, Inches(1))

    h = doc.add_heading("4.2 Planned Comparison with Existing Tools", level=1)
    for run in h.runs:
        run.font.size = Pt(14)

    intro = (
        "Existing handwriting supports for visually impaired learners—tactile aids, "
        "raised-line materials, tools such as LightWrite, braille-based systems, and "
        "multilingual handwriting apps—help, but often lack broad language coverage, "
        "adaptable interfaces, and meaningful feedback. The planned app targets those gaps "
        "with combined script support, voice guidance, and integrated review on a standard "
        "mobile device."
    )
    doc.add_paragraph(intro)

    doc.add_paragraph(
        "The table contrasts typical categories (not every product). Symbols summarise "
        "relative strength; see the legend below the table."
    )

    headers = [
        "Feature",
        "Tactile /\nraised-line",
        "Specialised\ndigital aids",
        "Braille-\nfirst",
        "Typical\napps",
        "Proposed\napp",
    ]

    # (label, tactile, spec_digital, braille, typical_apps, proposed)
    rows = [
        ("Multilingual letters & numerals", "mid", "mid", "mid", "mid", "ok"),
        ("No extra peripherals", "no", "mid", "mid", "ok", "ok"),
        ("Voice guidance while writing", "no", "mid", "mid", "no", "ok"),
        ("Automated feedback after writing", "no", "mid", "mid", "mid", "ok"),
        ("On-device ink (privacy)", "no", "mid", "mid", "mid", "ok"),
    ]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for j, title in enumerate(headers):
        _set_cell_text(hdr_cells[j], title.replace("\n", " "), bold=True)
        hdr_cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i, (label, a, b, c, d, e) in enumerate(rows, start=1):
        syms = [_SYM[a], _SYM[b], _SYM[c], _SYM[d], _SYM[e]]
        _set_cell_text(table.rows[i].cells[0], label, bold=True)
        for j, sym in enumerate(syms, start=1):
            ccell = table.rows[i].cells[j]
            _set_cell_text(ccell, sym, bold=False)
            ccell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    leg = doc.add_paragraph(_LEGEND)
    for run in leg.runs:
        run.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph()

    h2 = doc.add_heading("Narrative summary (paraphrased)", level=2)
    for run in h2.runs:
        run.font.size = Pt(13)

    bullets = [
        (
            "Multilingual writing support",
            "Unlike many tools centred on one script, the app is planned to cover English "
            "upper- and lowercase, Arabic numerals, and Chinese numerals for wider everyday use.",
        ),
        (
            "Independence and built-in feedback",
            "Instead of relying on specialised hardware or sighted help, the design emphasises "
            "solo practice on a phone or tablet with ongoing voice prompts and post-writing "
            "feedback—uncommon in current handwriting tools for this audience.",
        ),
    ]
    for bold_label, body in bullets:
        p = doc.add_paragraph(style="List Bullet")
        r0 = p.add_run(f"{bold_label}. ")
        r0.bold = True
        p.add_run(body)

    note = doc.add_paragraph()
    note.add_run("Note: ").bold = True
    note.add_run(
        "Columns describe categories; individual products differ. The “Proposed app” column "
        "matches the planned scope in this report."
    )

    return doc


def main() -> None:
    doc = build_document()
    out = _ROOT / "docs" / "Section_4.2_Comparison_with_Existing_Tools.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
