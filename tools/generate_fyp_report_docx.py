#!/usr/bin/env python3
"""Build a Word (.docx) with Sections 3.2-3.4 for pasting into the FYP report."""

import importlib.util
import sys
from io import BytesIO
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(_ROOT / ".doc_tools"))

from docx import Document  # noqa: E402
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402


def _load_pipeline_renderer():
    path = _ROOT / "tools" / "generate_section_3_3_3_4_pdf.py"
    spec = importlib.util.spec_from_file_location("pipeline_pdf", path)
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod.render_pipeline_figure_png


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(14 if level == 1 else 13)
    return None


def _add_bullet(doc: Document, bold_label: str, body: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r0 = p.add_run(f"{bold_label}: ")
    r0.bold = True
    p.add_run(body)


def build_document() -> Document:
    doc = Document()
    sect = doc.sections[0]
    sect.top_margin = Inches(1)
    sect.bottom_margin = Inches(1)
    sect.left_margin = Inches(1)
    sect.right_margin = Inches(1)

    doc.add_paragraph(
        "FYP report excerpts (Chapter 3) — generated for import into Microsoft Word."
    ).runs[0].italic = True

    # --- 3.2 ---
    _add_heading(doc, "3.2 Development Stack and Technical Framework", level=1)
    doc.add_paragraph(
        "Point-form summary of the stack; each item names the technique and its role in the build."
    )

    for title, body in [
        (
            "Figma",
            "High-fidelity prototyping and UI/UX design, including early checks of contrast and scaling "
            "before implementation.",
        ),
        (
            "Flutter (Dart)",
            "Primary cross-platform framework (single codebase for Android and iOS) with Material "
            "widgets suited to responsive, low-vision-friendly layouts.",
        ),
        (
            "Google ML Kit (Digital Ink Recognition)",
            "google_mlkit_digital_ink_recognition runs on-device: ink is recognised locally without "
            "sending stroke traces to a remote server.",
        ),
        (
            "Flutter APIs and flutter_tts",
            "Host TTS for real-time voice guidance and alignment with platform speech and accessibility "
            "ecosystems.",
        ),
        (
            "Bilingual interface and localization approach",
            "English/Chinese UI strings and Cantonese/English speech are driven by app state and "
            "saved preferences, not by intl or flutter_localizations.",
        ),
        (
            "shared_preferences, Semantics, and flutter_lints",
            "Persists UI and voice languages; Semantics supports screen readers; flutter_lints applies "
            "recommended Dart analysis rules.",
        ),
    ]:
        _add_bullet(doc, title, body)

    doc.add_page_break()

    # --- 3.3 opening ---
    _add_heading(doc, "3.3 App Components and Key Features", level=1)
    doc.add_paragraph(
        "This subsection describes interface structure, on-device ML Kit inference, and local "
        "storage, using the same terms as the codebase (Flutter/Dart, Digital Ink, flutter_tts, "
        "shared_preferences). Figure 1 outlines the practice-time pipeline."
    )

    render_png = _load_pipeline_renderer()
    img = render_png()
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = pic_para.add_run()
    # Wide figure for booklet / 2-up layouts; adjust in Word if margins differ.
    run.add_picture(buf, width=Inches(7.15))

    cap = doc.add_paragraph(
        "Figure 1. On-device practice pipeline (HC = high-contrast canvas; TTS = text-to-speech). "
        "Arrows: main control/data flow in one session."
    )
    cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(10)

    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run("Frontend (presentation layer)")
    r.bold = True

    for title, body in [
        (
            "High-contrast presentation and enlarged typography",
            "Dark scaffold, large base text sizes, and global linear text scaling aid low vision; a "
            "long-press on the canvas toggles a higher-contrast drawing view.",
        ),
        (
            "Navigation and task routing",
            "Home splits work into lowercase, uppercase, and numerals (grid then shared workspace). "
            "Poem Practice appears when the UI language is English (short phrases and longer verses).",
        ),
        (
            "Voice-guided writing workspace",
            "Template plus digital-ink canvas; TTS in Cantonese or English can differ from on-screen "
            "Chinese/English labels. Double-tap replays, pinch zooms, brush/eraser; glyph drills add "
            "round presets and a progress bar.",
        ),
        (
            "Passage-oriented (poem) composition mode",
            "Scrollable passage text; the learner writes one active grapheme at a time while speech "
            "tracks the current word and character.",
        ),
        (
            "Landscape layout and orientation control (poem practice only)",
            "Only poem practice uses a portrait/landscape split: portrait stacks word, progress, and "
            "canvas; landscape places a ~100px left strip (word, progress, circular speak/clear/next) "
            "beside a wide full-height canvas. Landscape applies on physical rotation or when the app "
            "bar pins orientation via SystemChrome (landscape left/right vs portrait-up); dispose "
            "restores portrait-up. Single-glyph practice has no such split.",
        ),
        (
            "Terminal analytics for passage tasks",
            "After the last token, TTS plus a Poem Summary dialog report overall accuracy, mean time "
            "per character, session duration, and difficult words.",
        ),
        (
            "Multilingual coverage",
            "Latin upper/lowercase, Arabic numerals, Traditional Chinese numerals (1-10 in the grid); "
            "bilingual UI with Cantonese or English TTS.",
        ),
        (
            "Accessibility metadata",
            "Semantics describe canvas actions (replay, high contrast, zoom) for platform screen "
            "readers.",
        ),
    ]:
        _add_bullet(doc, title, body)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Backend and local data handling (client-only)")
    r2.bold = True

    doc.add_paragraph(
        "Client-only design: no login, no cloud ink API, no Firebase. ML Kit scores strokes on-device. "
        "shared_preferences stores UI/voice locales and JSON lists of recent scores per practice key "
        "(not synced across devices). High-contrast canvas is session state, not a saved theme."
    )

    # --- 3.4 ---
    _add_heading(doc, "3.4 Core Features", level=1)

    for title, body in [
        (
            "(1) Voice-guided handwriting practice",
            "Spoken steps scaffold graphemes, words, and short lines. UI language and TTS language are "
            "independent (e.g. Chinese labels with English or Cantonese speech). On-device scoring "
            "drives spoken percentages and, in poem mode, auto-advance or retry prompts.",
        ),
        (
            "(2) Poem practice with continuous writing and structured feedback",
            "Same ML Kit path as drills, embedded in English passages; TTS follows the active span; "
            "scrollable text keeps context. End-of-session analytics combine speech and a summary UI "
            "(accuracy, timing, weak words).",
        ),
        (
            "(3) Multilingual script support",
            "Latin cases, Arabic digits, and Traditional Chinese numerals plus bilingual UI and "
            "flexible TTS; custom practice fonts and accessibility features keep tasks readable while "
            "all inference stays local.",
        ),
    ]:
        p = doc.add_paragraph()
        r0 = p.add_run(f"{title}. ")
        r0.bold = True
        p.add_run(body)

    return doc


def main() -> None:
    out = _ROOT / "docs" / "FYP_Report_Sections_3.2_to_3.4.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(out))
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
