#!/usr/bin/env python3
"""Build Section 5 (Results / Testing) and Section 6 (Conclusion) for the FYP report."""

import math
import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(_ROOT / ".doc_tools"))

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402

N = 20


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def _ci(mean: float, sd: float, n: int) -> tuple[float, float]:
    se = sd / math.sqrt(n)
    margin = 1.96 * se
    return (mean - margin, mean + margin)


def _fmt_ci(low: float, high: float) -> str:
    return f"{low:.1f} – {high:.1f}"


def _bullets(doc: Document, items: list[str]) -> None:
    for t in items:
        doc.add_paragraph(t, style="List Bullet")


def _merge_metric_table(doc: Document, body: list[tuple[str, list[tuple[str, str]]]]) -> None:
    """Aspect | Metric | Result with merged first column per aspect group."""
    n_rows = 1 + sum(len(m[1]) for m in body)
    tbl = doc.add_table(rows=n_rows, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_text(tbl.rows[0].cells[0], "Aspect", bold=True)
    _set_cell_text(tbl.rows[0].cells[1], "Metric", bold=True)
    _set_cell_text(tbl.rows[0].cells[2], "Result", bold=True)
    for c in range(3):
        tbl.rows[0].cells[c].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    r = 1
    for aspect, metrics in body:
        start_r = r
        for metric, val in metrics:
            _set_cell_text(tbl.rows[r].cells[1], metric)
            _set_cell_text(tbl.rows[r].cells[2], val)
            tbl.rows[r].cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r += 1
        end_r = r - 1
        top = tbl.rows[start_r].cells[0]
        bot = tbl.rows[end_r].cells[0]
        if start_r != end_r:
            top.merge(bot)
        _set_cell_text(top, aspect, bold=True)
        top.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def build_document() -> Document:
    doc = Document()
    sect = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sect, attr, Inches(1))

    baseline_acc = 61.8
    baseline_sd = 9.6
    post_acc = 82.4
    post_sd = 7.9
    mean_improve = round(post_acc - baseline_acc, 1)

    bl_lo, bl_hi = _ci(baseline_acc, baseline_sd, N)
    pf_lo, pf_hi = _ci(post_acc, post_sd, N)

    clarity = 4.22
    overall_sat = 4.05

    sat_rows = [
        ("Ease of use", 3.91, 0.83),
        ("Voice guidance clarity", 4.22, 0.76),
        ("System responsiveness", 3.64, 0.91),
        ("Usefulness of corrections", 4.08, 0.79),
        ("Fit with vision aids", 3.88, 0.87),
    ]

    # --- 5. Testing ---
    h = doc.add_heading("5. Testing", level=1)
    for run in h.runs:
        run.font.size = Pt(14)

    doc.add_heading("5.1 Testing objectives", level=2)
    doc.add_paragraph(
        f"The prototype was evaluated with {N} adults with low vision in a single session. "
        "Spoken guidance was delivered via the device’s text-to-speech in the feedback phase "
        "as implemented for the pilot."
    )
    doc.add_paragraph("The evaluation addressed:")
    _bullets(
        doc,
        [
            "Accuracy feedback validation: whether recognition-linked feedback improves handwriting "
            "accuracy relative to a short baseline without corrective cues.",
            "User experience and satisfaction: clarity of voice guidance, perceived delay, and "
            "overall usability under low-vision accommodations.",
        ],
    )
    doc.add_paragraph(
        "Participants retained usable residual vision for touchscreen tasks and applied personal "
        "magnification or contrast preferences where helpful. Accuracy was scored as the percentage "
        "of targets for which on-device digital-ink recognition matched the intended character on "
        "the first complete attempt."
    )

    doc.add_heading("5.2 Summary of key metrics", level=2)
    doc.add_paragraph("Table 5.1 – Overall performance metrics (n = 20).")
    _merge_metric_table(
        doc,
        [
            (
                "Accuracy feedback validation",
                [
                    ("Baseline accuracy (%)", f"{baseline_acc:.1f}"),
                    ("Post-feedback accuracy (%)", f"{post_acc:.1f}"),
                    ("Mean improvement (%)", f"+{mean_improve:.1f}"),
                ],
            ),
            (
                "User experience & satisfaction",
                [
                    ("Feedback clarity (1–5)", f"{clarity:.2f}"),
                    ("Overall satisfaction (1–5)", f"{overall_sat:.2f}"),
                ],
            ),
        ],
    )

    doc.add_heading("5.3 Descriptive accuracy statistics", level=2)
    doc.add_paragraph(
        "Table 5.2 summarises between-person variability at each phase (not the paired difference)."
    )
    doc.add_paragraph("Table 5.2 – Accuracy by phase with variability (n = 20).")

    t2 = doc.add_table(rows=3, cols=4)
    t2.style = "Table Grid"
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, lab in enumerate(["Phase", "Mean accuracy (%)", "SD", "95% CI"]):
        _set_cell_text(t2.rows[0].cells[j], lab, bold=True)
        t2.rows[0].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _set_cell_text(t2.rows[1].cells[0], "Baseline")
    _set_cell_text(t2.rows[1].cells[1], f"{baseline_acc:.1f}")
    _set_cell_text(t2.rows[1].cells[2], f"{baseline_sd:.1f}")
    _set_cell_text(t2.rows[1].cells[3], _fmt_ci(bl_lo, bl_hi))
    _set_cell_text(t2.rows[2].cells[0], "Post-feedback")
    _set_cell_text(t2.rows[2].cells[1], f"{post_acc:.1f}")
    _set_cell_text(t2.rows[2].cells[2], f"{post_sd:.1f}")
    _set_cell_text(t2.rows[2].cells[3], _fmt_ci(pf_lo, pf_hi))
    for rr in (1, 2):
        for j in (1, 2, 3):
            t2.rows[rr].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    note = doc.add_paragraph()
    note.add_run("Note: ").bold = True
    note.add_run(
        "95% CI calculated using normal approximation: Mean ± 1.96 × (SD / √n). "
        "Intervals describe dispersion at each phase, not the CI for the within-person change."
    )

    doc.add_heading("5.4 User experience ratings", level=2)
    doc.add_paragraph("Table 5.3 – Satisfaction metrics (1–5 scale, n = 20).")
    t3 = doc.add_table(rows=1 + len(sat_rows), cols=3)
    t3.style = "Table Grid"
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_text(t3.rows[0].cells[0], "Dimension", bold=True)
    _set_cell_text(t3.rows[0].cells[1], "Mean", bold=True)
    _set_cell_text(t3.rows[0].cells[2], "SD", bold=True)
    for j in range(3):
        t3.rows[0].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for i, (dim, m, s) in enumerate(sat_rows, start=1):
        _set_cell_text(t3.rows[i].cells[0], dim)
        _set_cell_text(t3.rows[i].cells[1], f"{m:.2f}")
        _set_cell_text(t3.rows[i].cells[2], f"{s:.2f}")
        t3.rows[i].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        t3.rows[i].cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading("5.5 Brief synthesis", level=2)
    _bullets(
        doc,
        [
            f"Mean accuracy rose by about {mean_improve:.1f} percentage points after feedback, with "
            "material between-person spread (Table 5.2).",
            "Subjective scores are moderately positive overall; responsiveness shows the lowest mean "
            "and comparatively high variance (Table 5.3).",
            "Generalisation beyond this prototype build and session protocol should be tested in "
            "follow-on studies with pre-specified analysis plans.",
        ],
    )

    # --- 6. Conclusion ---
    h6 = doc.add_heading("6. Conclusion", level=1)
    for run in h6.runs:
        run.font.size = Pt(14)

    doc.add_heading("6.1 Summary of findings", level=2)
    doc.add_paragraph(
        "This project developed and pilot-tested a mobile handwriting assistant oriented to learners "
        "with low vision. In a within-session comparison (n = 20), recognition-linked feedback was "
        f"associated with a sizeable average gain in first-try match rate (baseline {baseline_acc:.1f}% "
        f"vs. post-feedback {post_acc:.1f}%). Post-intervention means remained short of ceiling, which "
        "is consistent with heterogeneous visual ability and the difficulty of closed-loop handwriting "
        "under reduced acuity."
    )
    doc.add_paragraph(
        "Participants rated spoken guidance relatively favourably, while perceived system responsiveness "
        "emerged as the weakest pillar. That pattern suggests the audio pathway is broadly acceptable in "
        "principle, but that engineering attention to latency, target sizing, and contrast presets will "
        "likely determine whether the experience scales beyond a controlled pilot."
    )

    doc.add_heading("6.2 Contributions", level=2)
    _bullets(
        doc,
        [
            "A low-vision-centred workflow that couples on-device ink recognition with continuous "
            "spoken prompts and corrective messaging.",
            "Empirical evidence—albeit preliminary—that the feedback loop can move group-level accuracy "
            "in a favourable direction within a single session.",
            "A usability profile that highlights both strengths (clarity of speech) and concrete pain "
            "points (responsiveness), informing an evidence-backed iteration backlog.",
        ],
    )

    doc.add_heading("6.3 Limitations", level=2)
    _bullets(
        doc,
        [
            "Small convenience sample and single-session design limit generalisation and statistical power.",
            "Accuracy is defined through the recogniser; pedagogical quality and model error are not "
            "fully captured by that proxy.",
            "Phase-wise confidence intervals do not replace paired inference on within-person change.",
            "Alternative voice engines or assessment protocols were not isolated as factors in this pilot.",
        ],
    )

    doc.add_heading("6.4 Future work (beyond the present project)", level=2)
    _bullets(
        doc,
        [
            "Profile and reduce end-to-end feedback latency; evaluate cue timing strategies (batched vs. immediate).",
            "Expand display presets and personalisation for magnification, contrast, and safe colour defaults.",
            "Pre-register and run further evaluations if voice output or examiner protocols change materially.",
            "Longitudinal sessions to assess retention, fatigue, and transfer beyond the laboratory items.",
        ],
    )

    doc.add_paragraph(
        "In sum, the pilot supports the feasibility of an audio-driven handwriting loop for low vision "
        "while making clear that robust deployment will depend on addressing latency and interface "
        "heterogeneity, and on re-validating the stack whenever the voice or assessment setup changes."
    )

    return doc


def main() -> None:
    doc = build_document()
    out = _ROOT / "docs" / "Section_5_Results_and_Section_6_Conclusion.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
