#!/usr/bin/env python3
"""End-of-term report (zh-TW): §5–§7; strings from tools/_zh_blocks.json."""

from __future__ import annotations

import json
import math
import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(_ROOT / ".doc_tools"))

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402

N = 20
_ZH_PATH = Path(__file__).resolve().parent / "_zh_blocks.json"


def _load_zh() -> dict:
    return json.loads(_ZH_PATH.read_text(encoding="utf-8"))


def _h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.size = Pt(14 if level == 1 else 12)


def _set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    r = cell.paragraphs[0].add_run(text)
    r.bold = bold
    r.font.size = Pt(10)


def _ci(m: float, sd: float, n: int) -> tuple[float, float]:
    se = sd / math.sqrt(n)
    mar = 1.96 * se
    return (m - mar, m + mar)


def _fmt_ci(lo: float, hi: float) -> str:
    return f"{lo:.1f} – {hi:.1f}"


def _bullets(doc: Document, items: list[str]) -> None:
    for t in items:
        doc.add_paragraph(t, style="List Bullet")


def _merge_table_zh(
    doc: Document,
    hdr: list[str],
    groups: list[tuple[str, list[tuple[str, str]]]],
) -> None:
    rows = 1 + sum(len(g[1]) for g in groups)
    t = doc.add_table(rows=rows, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell(t.rows[0].cells[0], hdr[0], True)
    _set_cell(t.rows[0].cells[1], hdr[1], True)
    _set_cell(t.rows[0].cells[2], hdr[2], True)
    for c in range(3):
        t.rows[0].cells[c].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = 1
    for aspect, metrics in groups:
        sr = r
        for metric, val in metrics:
            _set_cell(t.rows[r].cells[1], metric)
            _set_cell(t.rows[r].cells[2], val)
            t.rows[r].cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r += 1
        top, bot = t.rows[sr].cells[0], t.rows[r - 1].cells[0]
        if sr != r - 1:
            top.merge(bot)
        _set_cell(top, aspect, True)
        top.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def build_document() -> Document:
    Z = _load_zh()
    doc = Document()
    s = doc.sections[0]
    for a in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(s, a, Inches(1))

    base_m, base_sd = 62.1, 9.4
    post_m, post_sd = 81.9, 8.1
    delta = round(post_m - base_m, 1)
    bl_lo, bl_hi = _ci(base_m, base_sd, N)
    pf_lo, pf_hi = _ci(post_m, post_sd, N)

    clarity = 4.18
    overall = 4.02

    sat = [
        (Z["sat_zh"][0], 3.95, 0.81),
        (Z["sat_zh"][1], 4.18, 0.74),
        (Z["sat_zh"][2], 3.61, 0.89),
        (Z["sat_zh"][3], 4.05, 0.77),
        (Z["sat_zh"][4], 3.86, 0.85),
        (Z["sat_zh"][5], 3.79, 0.88),
    ]

    intro = doc.add_paragraph(Z["intro"].format(n=N))
    intro.runs[0].italic = True

    _bullets(doc, list(Z["road"]))

    _h(doc, Z["h5"], 1)

    doc.add_heading(Z["h51t"], level=2)
    doc.add_paragraph(Z["h51p"])
    _bullets(doc, list(Z["h51b"]))

    doc.add_heading(Z["h52t"], level=2)
    doc.add_paragraph(Z["t51cap"])

    snap = list(zip(Z["snap_k"], [d.format(n=N) if "{n}" in d else d for d in Z["snap_d"]]))
    t0 = doc.add_table(rows=1 + len(snap), cols=2)
    t0.style = "Table Grid"
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell(t0.rows[0].cells[0], Z["item"], True)
    _set_cell(t0.rows[0].cells[1], Z["desc"], True)
    for i, (k, v) in enumerate(snap, start=1):
        _set_cell(t0.rows[i].cells[0], k, True)
        _set_cell(t0.rows[i].cells[1], v)
        t0.rows[i].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        t0.rows[i].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    t0.rows[0].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    t0.rows[0].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_heading(Z["h53t"], level=2)
    doc.add_paragraph(Z["t52cap"])
    _merge_table_zh(
        doc,
        Z["merge_hdr"],
        [
            (
                Z["acc"],
                [
                    (Z["mbase"], f"{base_m:.1f}"),
                    (Z["mpost"], f"{post_m:.1f}"),
                    (Z["mgain"], f"+{delta:.1f}"),
                ],
            ),
            (
                Z["glob"],
                [
                    (Z["mclar"], f"{clarity:.2f}"),
                    (Z["msat"], f"{overall:.2f}"),
                ],
            ),
        ],
    )

    doc.add_paragraph(Z["t53cap"])
    cols = Z["phase_cols"]
    t2 = doc.add_table(rows=3, cols=4)
    t2.style = "Table Grid"
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, lab in enumerate(cols):
        _set_cell(t2.rows[0].cells[j], lab, True)
        t2.rows[0].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _set_cell(t2.rows[1].cells[0], Z["ph_base"])
    _set_cell(t2.rows[1].cells[1], f"{base_m:.1f}")
    _set_cell(t2.rows[1].cells[2], f"{base_sd:.1f}")
    _set_cell(t2.rows[1].cells[3], _fmt_ci(bl_lo, bl_hi))
    _set_cell(t2.rows[2].cells[0], Z["ph_post"])
    _set_cell(t2.rows[2].cells[1], f"{post_m:.1f}")
    _set_cell(t2.rows[2].cells[2], f"{post_sd:.1f}")
    _set_cell(t2.rows[2].cells[3], _fmt_ci(pf_lo, pf_hi))
    for rr in (1, 2):
        for j in (1, 2, 3):
            t2.rows[rr].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    n = doc.add_paragraph()
    note = Z["note"]
    if note.startswith("\u8a3b\u3002"):
        n.add_run("\u8a3b\u3002").bold = True
        n.add_run(note[2:])
    else:
        n.add_run(note)

    doc.add_heading(Z["h54t"], level=2)
    doc.add_paragraph(Z["t54cap"])
    t3 = doc.add_table(rows=1 + len(sat), cols=3)
    t3.style = "Table Grid"
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell(t3.rows[0].cells[0], Z["dim"], True)
    _set_cell(t3.rows[0].cells[1], Z["mean"], True)
    _set_cell(t3.rows[0].cells[2], "SD", True)
    for j in range(3):
        t3.rows[0].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for i, (d, m, sd) in enumerate(sat, start=1):
        _set_cell(t3.rows[i].cells[0], d)
        _set_cell(t3.rows[i].cells[1], f"{m:.2f}")
        _set_cell(t3.rows[i].cells[2], f"{sd:.2f}")
        t3.rows[i].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        t3.rows[i].cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(Z["t55cap"])
    ranked = sorted(sat, key=lambda x: x[1])
    t4 = doc.add_table(rows=1 + len(ranked), cols=3)
    t4.style = "Table Grid"
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell(t4.rows[0].cells[0], Z["col_rank"], True)
    _set_cell(t4.rows[0].cells[1], Z["dim"], True)
    _set_cell(t4.rows[0].cells[2], Z["mean"], True)
    for j in range(3):
        t4.rows[0].cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for i, (d, m, sd) in enumerate(ranked, start=1):
        _set_cell(t4.rows[i].cells[0], str(i))
        _set_cell(t4.rows[i].cells[1], d)
        _set_cell(t4.rows[i].cells[2], f"{m:.2f}")
        t4.rows[i].cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading(Z["h55t"], level=2)
    h55 = [Z["h55b"][0].format(delta=delta), Z["h55b"][1], Z["h55b"][2]]
    _bullets(doc, h55)

    _h(doc, Z["h6"], 1)
    doc.add_paragraph(Z["h6i"])

    doc.add_heading(Z["h61t"], level=2)
    doc.add_paragraph(Z["h61p1"].format(delta=delta, post_m=f"{post_m:.0f}"))
    doc.add_paragraph(Z["h61p2"])

    doc.add_heading(Z["h62t"], level=2)
    doc.add_paragraph(Z["h62p"])

    doc.add_heading(Z["h63t"], level=2)
    doc.add_paragraph(Z["h63p"])

    doc.add_heading(Z["h64t"], level=2)
    doc.add_paragraph(Z["h64p"])

    doc.add_heading(Z["h65t"], level=2)
    doc.add_paragraph(Z["h65p"])

    doc.add_heading(Z["h66t"], level=2)
    h66 = [
        Z["h66b"][0].format(base_m=f"{base_m:.1f}", post_m=f"{post_m:.1f}"),
        Z["h66b"][1],
        Z["h66b"][2],
    ]
    _bullets(doc, h66)

    _h(doc, Z["h7"], 1)
    doc.add_paragraph(Z["h7i"])

    doc.add_heading(Z["h71t"], level=2)
    _bullets(doc, list(Z["h71b"]))

    doc.add_heading(Z["h72t"], level=2)
    _bullets(doc, list(Z["h72b"]))

    doc.add_heading(Z["h73t"], level=2)
    _bullets(doc, list(Z["h73b"]))

    doc.add_paragraph(Z["close"])

    return doc


def main() -> None:
    doc = build_document()
    out = _ROOT / "docs" / "End_of_Term_Report_zh_TW.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
